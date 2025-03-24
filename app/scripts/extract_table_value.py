import sys
import argparse
from docx import Document

def extract_table_value(docx_path):
    doc = Document(docx_path)
    if len(doc.tables) < 3:
        sys.exit("No se encontró la tercera tabla.")

    table = doc.tables[2]

    if len(table.rows) < 2:
        sys.exit("No se encontró la segunda fila en la tercera tabla.")

    row = table.rows[1]  # Segunda fila

    if len(row.cells) < 3:
        sys.exit("No se encontró la tercera columna en la segunda fila.")

    return row.cells[2].text

def main():
    parser = argparse.ArgumentParser(
        description="Extrae el contenido de la tercera tabla, segunda fila y tercera columna de un DOCX."
    )
    parser.add_argument("--docx", required=True, help="Ruta al archivo DOCX.")
    args = parser.parse_args()

    value = extract_table_value(args.docx)
    print(value)

if __name__ == "__main__":
    main()
